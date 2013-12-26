# encoding: utf-8

"""
Test suite for the docx.parts module
"""

from __future__ import absolute_import, print_function, unicode_literals

import pytest

from mock import Mock

from docx.enum.shape import WD_INLINE_SHAPE
from docx.oxml.parts import CT_Document
from docx.oxml.shared import nsmap
from docx.parts import _Body, _Document, InlineShape, InlineShapes
from docx.table import Table
from docx.text import Paragraph

from .oxml.unitdata.dml import (
    a_blip, a_blipFill, a_drawing, a_graphic, a_graphicData, a_pic, an_inline
)
from .oxml.unitdata.parts import a_body, a_document
from .oxml.unitdata.table import (
    a_gridCol, a_tbl, a_tblGrid, a_tblPr, a_tc, a_tr
)
from .oxml.unitdata.text import a_p, a_sectPr, an_r
from .unitutil import (
    function_mock, class_mock, initializer_mock, instance_mock
)


class Describe_Document(object):

    def it_can_be_constructed_by_opc_part_factory(
            self, oxml_fromstring_, init):
        # mockery ----------------------
        partname, content_type, blob, document_elm, package = (
            Mock(name='partname'), Mock(name='content_type'),
            Mock(name='blob'), Mock(name='document_elm'),
            Mock(name='package')
        )
        oxml_fromstring_.return_value = document_elm
        # exercise ---------------------
        doc = _Document.load(partname, content_type, blob, package)
        # verify -----------------------
        oxml_fromstring_.assert_called_once_with(blob)
        init.assert_called_once_with(
            partname, content_type, document_elm, package
        )
        assert isinstance(doc, _Document)

    def it_has_a_body(self, document_body_fixture):
        document, _Body_, body_elm = document_body_fixture
        _body = document.body
        _Body_.assert_called_once_with(body_elm)
        assert _body is _Body_.return_value

    def it_can_serialize_to_xml(self, document_blob_fixture):
        document, document_elm, serialize_part_xml_ = document_blob_fixture
        blob = document.blob
        serialize_part_xml_.assert_called_once_with(document_elm)
        assert blob is serialize_part_xml_.return_value

    def it_provides_access_to_the_inline_shapes_in_the_document(
            self, inline_shapes_fixture):
        document, InlineShapes_, body_elm = inline_shapes_fixture
        inline_shapes = document.inline_shapes
        InlineShapes_.assert_called_once_with(body_elm)
        assert inline_shapes is InlineShapes_.return_value

    # fixtures -------------------------------------------------------

    @pytest.fixture
    def document_blob_fixture(self, request, serialize_part_xml_):
        document_elm = instance_mock(request, CT_Document)
        document = _Document(None, None, document_elm, None)
        return document, document_elm, serialize_part_xml_

    @pytest.fixture
    def document_body_fixture(self, request, _Body_):
        document_elm = (
            a_document().with_nsdecls().with_child(
                a_body())
        ).element
        body_elm = document_elm[0]
        document = _Document(None, None, document_elm, None)
        return document, _Body_, body_elm

    @pytest.fixture
    def _Body_(self, request):
        return class_mock(request, 'docx.parts._Body')

    @pytest.fixture
    def init(self, request):
        return initializer_mock(request, _Document)

    @pytest.fixture
    def InlineShapes_(self, request):
        return class_mock(request, 'docx.parts.InlineShapes')

    @pytest.fixture
    def inline_shapes_fixture(self, request, InlineShapes_):
        document_elm = (
            a_document().with_nsdecls().with_child(
                a_body())
        ).element
        body_elm = document_elm[0]
        document = _Document(None, None, document_elm, None)
        return document, InlineShapes_, body_elm

    @pytest.fixture
    def oxml_fromstring_(self, request):
        return function_mock(request, 'docx.parts.oxml_fromstring')

    @pytest.fixture
    def serialize_part_xml_(self, request):
        return function_mock(request, 'docx.parts.serialize_part_xml')


class Describe_Body(object):

    def it_can_add_a_paragraph(self, add_paragraph_fixture):
        body, expected_xml = add_paragraph_fixture
        p = body.add_paragraph()
        assert body._body.xml == expected_xml
        assert isinstance(p, Paragraph)

    def it_can_add_a_table(self, add_table_fixture):
        body, expected_xml = add_table_fixture
        table = body.add_table(rows=1, cols=1)
        assert body._body.xml == expected_xml
        assert isinstance(table, Table)

    def it_can_clear_itself_of_all_content_it_holds(
            self, clear_content_fixture):
        body, expected_xml = clear_content_fixture
        _body = body.clear_content()
        assert body._body.xml == expected_xml
        assert _body is body

    def it_provides_access_to_the_paragraphs_it_contains(
            self, body_with_paragraphs):
        body = body_with_paragraphs
        paragraphs = body.paragraphs
        assert len(paragraphs) == 2
        for p in paragraphs:
            assert isinstance(p, Paragraph)

    def it_provides_access_to_the_tables_it_contains(
            self, body_with_tables):
        body = body_with_tables
        tables = body.tables
        assert len(tables) == 2
        for table in tables:
            assert isinstance(table, Table)

    # fixtures -------------------------------------------------------

    @pytest.fixture(params=[
        (0, False), (1, False), (0, True), (1, True)
    ])
    def add_paragraph_fixture(self, request):
        p_count, has_sectPr = request.param
        # body element -----------------
        body_bldr = self._body_bldr(p_count=p_count, sectPr=has_sectPr)
        body_elm = body_bldr.element
        body = _Body(body_elm)
        # expected XML -----------------
        p_count += 1
        body_bldr = self._body_bldr(p_count=p_count, sectPr=has_sectPr)
        expected_xml = body_bldr.xml()
        return body, expected_xml

    @pytest.fixture(params=[(0, False), (0, True), (1, False), (1, True)])
    def add_table_fixture(self, request):
        p_count, has_sectPr = request.param
        body_bldr = self._body_bldr(p_count=p_count, sectPr=has_sectPr)
        body = _Body(body_bldr.element)

        tbl_bldr = self._tbl_bldr()
        body_bldr = self._body_bldr(
            p_count=p_count, tbl_bldr=tbl_bldr, sectPr=has_sectPr
        )
        expected_xml = body_bldr.xml()

        return body, expected_xml

    @pytest.fixture
    def body_with_paragraphs(self):
        body_elm = (
            a_body().with_nsdecls()
                    .with_child(a_p())
                    .with_child(a_p())
                    .element
        )
        return _Body(body_elm)

    @pytest.fixture
    def body_with_tables(self):
        body_elm = (
            a_body().with_nsdecls()
                    .with_child(a_tbl())
                    .with_child(a_tbl())
                    .element
        )
        return _Body(body_elm)

    @pytest.fixture(params=[False, True])
    def clear_content_fixture(self, request):
        has_sectPr = request.param
        # body element -----------------
        body_bldr = a_body().with_nsdecls()
        body_bldr.with_child(a_p())
        if has_sectPr:
            body_bldr.with_child(a_sectPr())
        body_elm = body_bldr.element
        body = _Body(body_elm)
        # expected XML -----------------
        body_bldr = a_body().with_nsdecls()
        if has_sectPr:
            body_bldr.with_child(a_sectPr())
        expected_xml = body_bldr.xml()
        return body, expected_xml

    def _body_bldr(self, p_count=0, tbl_bldr=None, sectPr=False):
        body_bldr = a_body().with_nsdecls()
        for i in range(p_count):
            body_bldr.with_child(a_p())
        if tbl_bldr is not None:
            body_bldr.with_child(tbl_bldr)
        if sectPr:
            body_bldr.with_child(a_sectPr())
        return body_bldr

    def _tbl_bldr(self, rows=1, cols=1):
        tblPr_bldr = a_tblPr()

        tblGrid_bldr = a_tblGrid()
        for i in range(cols):
            tblGrid_bldr.with_child(a_gridCol())

        tbl_bldr = a_tbl()
        tbl_bldr.with_child(tblPr_bldr)
        tbl_bldr.with_child(tblGrid_bldr)
        for i in range(rows):
            tr_bldr = self._tr_bldr(cols)
            tbl_bldr.with_child(tr_bldr)

        return tbl_bldr

    def _tc_bldr(self):
        return a_tc().with_child(a_p())

    def _tr_bldr(self, cols):
        tr_bldr = a_tr()
        for i in range(cols):
            tc_bldr = self._tc_bldr()
            tr_bldr.with_child(tc_bldr)
        return tr_bldr


class DescribeInlineShape(object):

    def it_knows_what_type_of_shape_it_is(self, shape_type_fixture):
        inline_shape, inline_shape_type = shape_type_fixture
        assert inline_shape.type == inline_shape_type

    # fixtures -------------------------------------------------------

    @pytest.fixture(params=[
        'embed pic', 'link pic', 'link+embed pic', 'chart', 'smart art',
        'not implemented'
    ])
    def shape_type_fixture(self, request):
        if request.param == 'embed pic':
            inline = self._inline_with_picture(embed=True)
            shape_type = WD_INLINE_SHAPE.PICTURE

        elif request.param == 'link pic':
            inline = self._inline_with_picture(link=True)
            shape_type = WD_INLINE_SHAPE.LINKED_PICTURE

        elif request.param == 'link+embed pic':
            inline = self._inline_with_picture(embed=True, link=True)
            shape_type = WD_INLINE_SHAPE.LINKED_PICTURE

        elif request.param == 'chart':
            inline = self._inline_with_uri(nsmap['c'])
            shape_type = WD_INLINE_SHAPE.CHART

        elif request.param == 'smart art':
            inline = self._inline_with_uri(nsmap['dgm'])
            shape_type = WD_INLINE_SHAPE.SMART_ART

        elif request.param == 'not implemented':
            inline = self._inline_with_uri('foobar')
            shape_type = WD_INLINE_SHAPE.NOT_IMPLEMENTED

        return InlineShape(inline), shape_type

    def _inline_with_picture(self, embed=False, link=False):
        picture_ns = nsmap['pic']

        blip_bldr = a_blip()
        if embed:
            blip_bldr.with_embed('rId1')
        if link:
            blip_bldr.with_link('rId2')

        inline = (
            an_inline().with_nsdecls('wp', 'r').with_child(
                a_graphic().with_nsdecls().with_child(
                    a_graphicData().with_uri(picture_ns).with_child(
                        a_pic().with_nsdecls().with_child(
                            a_blipFill().with_child(
                                blip_bldr)))))
        ).element
        return inline

    def _inline_with_uri(self, uri):
        inline = (
            an_inline().with_nsdecls('wp').with_child(
                a_graphic().with_nsdecls().with_child(
                    a_graphicData().with_uri(uri)))
        ).element
        return inline


class DescribeInlineShapes(object):

    def it_knows_how_many_inline_shapes_it_contains(
            self, inline_shapes_fixture):
        inline_shapes, inline_shape_count = inline_shapes_fixture
        assert len(inline_shapes) == inline_shape_count

    def it_can_iterate_over_its_InlineShape_instances(
            self, inline_shapes_fixture):
        inline_shapes, inline_shape_count = inline_shapes_fixture
        actual_count = 0
        for inline_shape in inline_shapes:
            assert isinstance(inline_shape, InlineShape)
            actual_count += 1
        assert actual_count == inline_shape_count

    def it_provides_indexed_access_to_inline_shapes(
            self, inline_shapes_fixture):
        inline_shapes, inline_shape_count = inline_shapes_fixture
        for idx in range(-inline_shape_count, inline_shape_count):
            inline_shape = inline_shapes[idx]
            assert isinstance(inline_shape, InlineShape)

    def it_raises_on_indexed_access_out_of_range(
            self, inline_shapes_fixture):
        inline_shapes, inline_shape_count = inline_shapes_fixture
        with pytest.raises(IndexError):
            too_low = -1 - inline_shape_count
            inline_shapes[too_low]
        with pytest.raises(IndexError):
            too_high = inline_shape_count
            inline_shapes[too_high]

    # fixtures -------------------------------------------------------

    @pytest.fixture
    def inline_shapes_fixture(self):
        inline_shape_count = 2
        body = (
            a_body().with_nsdecls('w', 'wp').with_child(
                a_p().with_child(
                    an_r().with_child(
                        a_drawing().with_child(
                            an_inline()))).with_child(
                    an_r().with_child(
                        a_drawing().with_child(
                            an_inline())
                    )
                )
            )
        ).element
        inline_shapes = InlineShapes(body)
        return inline_shapes, inline_shape_count